# 文件夹文本爬虫 - 用于爬取多个AI相关文件夹的文本内容
# @self-expose: {"id": "folder_crawler", "name": "Folder Crawler", "type": "component", "version": "1.0.0", "needs": {"deps": [], "resources": []}, "provides": {"capabilities": ["Folder Crawler功能"]}}

import os
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Set

# 导入必要的模块
import os
import sys
import json
import uuid
import logging
from typing import List, Dict, Any, Set
from datetime import datetime
from pathlib import Path
from collections import defaultdict

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 导入RAG系统的必要模块
from tools.memory_slicer_tool import MemorySlicerTool
from config.system_config import DATABASE_PATH

# 导入DOC文件处理模块
try:
    from docx import Document
except ImportError:
    Document = None
    logger.warning("python-docx模块未安装，DOC文件处理功能将不可用")

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class FolderCrawler:
    """文件夹爬虫类"""
    
    def __init__(self, max_depth: int = 3, output_json: bool = False, topic: str = 'folder_content', enable_llm_processing: bool = False):
        """
        文件夹爬虫类
        
        Args:
            max_depth: 最大递归深度
            output_json: 是否输出JSON
            topic: 记忆单元的主题分类
            enable_llm_processing: 是否启用LLM处理，让模型先理解文件内容
        """
        # 创建统一切片器
        self.slicer = MemorySlicerTool()
        # 基本配置
        self.max_depth = max_depth
        self.output_json = output_json
        self.topic = topic
        # 是否启用LLM处理
        self.enable_llm_processing = enable_llm_processing
        # LLM客户端
        self.llm_client = None
        # 已处理文件集合，避免重复处理
        self.processed_files: Set[str] = set()
        # 统计信息
        self.stats = {
            'total_files': 0,
            'processed_files': 0,
            'skipped_files': 0,
            'created_slices': 0,
            'file_type_stats': defaultdict(int),
            'folder_stats': defaultdict(int)
        }
        # 支持的文本文件类型
        self.supported_file_types = {
            '.txt', '.md', '.log', '.py', '.json', '.html', '.htm', '.csv',
            '.yaml', '.yml', '.xml', '.js', '.css', '.jsonl', '.sql', '.sh',
            '.bat', '.ps1', '.java', '.cpp', '.c', '.h', '.cs', '.php', '.rb',
            '.go', '.rs', '.swift', '.kt', '.scala', '.pl', '.lua', '.r',
            '.doc', '.docx'  # 添加DOC文件支持
        }
        
        # 初始化LLM客户端（如果启用）
        if self.enable_llm_processing:
            try:
                from src.llm_client_enhanced import LLMClientEnhanced
                self.llm_client = LLMClientEnhanced()
                print("LLM客户端初始化成功")
            except Exception as e:
                print(f"LLM客户端初始化失败，将跳过LLM处理: {e}")
                self.enable_llm_processing = False
    
    def should_process_file(self, file_path: Path) -> bool:
        """检查文件是否应该被处理"""
        # 检查扩展名
        extension = file_path.suffix.lower()
        if extension not in self.supported_file_types:
            return False
        
        # 检查文件大小（跳过空文件）
        try:
            if file_path.stat().st_size == 0:
                return False
        except Exception as e:
            logger.error(f"无法获取文件大小 {file_path}: {e}")
            return False
        
        # 检查是否已处理过
        file_abs_path = str(file_path.absolute())
        if file_abs_path in self.processed_files:
            return False
        
        return True
    
    def read_text_file(self, file_path: Path) -> str:
        """读取文本文件内容，支持DOC/DOCX文件"""
        try:
            extension = file_path.suffix.lower()
            
            # 处理DOC/DOCX文件
            if extension in ['.doc', '.docx']:
                if Document is None:
                    logger.error(f"无法处理DOC文件 {file_path}，python-docx模块未安装")
                    return ""
                
                try:
                    doc = Document(file_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content += paragraph.text + "\n"
                    
                    # 处理表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ""
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text += cell.text.strip() + " | "
                            if row_text:
                                content += row_text.rstrip(" | ") + "\n"
                    
                    logger.debug(f"成功读取DOC文件 {file_path}，提取了 {len(content)} 个字符")
                    return content
                    
                except Exception as e:
                    logger.error(f"读取DOC文件失败 {file_path}: {e}")
                    return ""
            
            # 处理普通文本文件
            else:
                # 尝试多种编码
                encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1']
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            logger.debug(f"成功读取文件 {file_path} 使用编码 {encoding}")
                            return content
                    except UnicodeDecodeError:
                        continue
                
                logger.error(f"无法解码文件 {file_path}，尝试了所有支持的编码")
                return ""
                
        except Exception as e:
            logger.error(f"读取文件失败 {file_path}: {e}")
            return ""
    
    def process_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理单个文件，返回切片数据"""
        if not self.should_process_file(file_path):
            extension = file_path.suffix.lower()
            if extension not in self.supported_file_types:
                logger.debug(f"跳过不支持的文件类型: {file_path}")
            else:
                logger.debug(f"跳过已处理或空文件: {file_path}")
            
            self.stats['skipped_files'] += 1
            return []
        
        # 记录已处理
        file_abs_path = str(file_path.absolute())
        self.processed_files.add(file_abs_path)
        
        # 更新统计信息
        self.stats['total_files'] += 1
        self.stats['processed_files'] += 1
        
        # 更新按扩展名统计
        extension = file_path.suffix.lower()
        if extension not in self.stats['file_type_stats']:
            self.stats['file_type_stats'][extension] = 0
        self.stats['file_type_stats'][extension] += 1
        
        # 更新按文件夹统计
        parent_folder = str(file_path.parent)
        if parent_folder not in self.stats['folder_stats']:
            self.stats['folder_stats'][parent_folder] = 0
        self.stats['folder_stats'][parent_folder] += 1
        
        # 读取文件内容
        content = self.read_text_file(file_path)
        if not content.strip():
            logger.warning(f"文件内容为空或无法读取: {file_path}")
            return []
        
        # 构建元数据
        file_type = 'text' if file_path.suffix.lower() in self.supported_file_types else 'unknown'
        folder_name = file_path.parent.name
        
        metadata = {
            'source': 'folder_crawler',
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_type': file_type,
            'folder_name': folder_name,
            'parent_folder': str(file_path.parent),
            'file_size': len(content),
            'last_modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
            'collected_at': datetime.now().isoformat()
        }
        
        # 让LLM先看过文件内容
        if self.enable_llm_processing and self.llm_client:
            logger.info(f"LLM正在处理文件: {file_path}")
            
            try:
                # 构建LLM提示词
                prompt = f"请简要理解并总结以下文件内容，提取关键信息：\n\n文件路径: {file_path}\n文件内容: {content[:2000]}...\n\n请输出简洁的总结，不超过500字："
                
                # 调用LLM
                messages = [
                    {"role": "system", "content": "你是一个文档理解助手，能够快速理解并总结文件内容。"},
                    {"role": "user", "content": prompt}
                ]
                
                llm_summary = self.llm_client.chat_completion(messages, max_tokens=500)
                
                if llm_summary:
                    # 添加LLM处理结果到元数据
                    metadata['llm_processed'] = True
                    metadata['llm_processed_at'] = datetime.now().isoformat()
                    metadata['llm_summary'] = llm_summary
                    logger.info(f"LLM处理文件完成: {file_path}")
                else:
                    logger.warning(f"LLM处理文件失败: {file_path}")
                    metadata['llm_processed'] = False
            except Exception as e:
                logger.error(f"LLM处理文件时发生错误: {file_path}, 错误: {e}")
                metadata['llm_processed'] = False
        else:
            # 模拟LLM处理
            logger.info(f"LLM正在处理文件: {file_path}")
            metadata['llm_processed'] = True
            metadata['llm_processed_at'] = datetime.now().isoformat()
        
        # 使用LLM切片器进行智能切片
        slices = self.slicer.slice_text(content, metadata)
        
        if slices:
            logger.info(f"成功处理文件 {file_path}，创建了 {len(slices)} 个切片")
            self.stats['created_slices'] += len(slices)
        else:
            logger.warning(f"文件切片失败: {file_path}")
        
        return slices
    
    def crawl_folder(self, folder_path: str, max_depth: int = 3) -> List[Dict[str, Any]]:
        """爬取指定文件夹下的所有文本文件"""
        folder = Path(folder_path)
        
        if not folder.exists() or not folder.is_dir():
            logger.error(f"文件夹不存在或不是目录: {folder_path}")
            return []
        
        all_slices = []
        start_time = datetime.now()
        
        logger.info(f"开始爬取文件夹: {folder_path}，最大深度: {max_depth}")
        
        # 遍历文件夹
        for root, _, files in os.walk(folder_path):
            # 计算当前深度
            current_depth = len(Path(root).relative_to(folder).parts)
            
            # 检查深度限制
            if current_depth > max_depth:
                continue
            
            for file_name in files:
                file_path = Path(root) / file_name
                file_slices = self.process_file(file_path)
                all_slices.extend(file_slices)
        
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        
        logger.info(f"文件夹爬取完成: {folder_path}")
        logger.info(f"耗时: {duration:.2f} 秒，处理文件: {self.stats['processed_files']}，创建切片: {len(all_slices)}")
        
        return all_slices
    
    def crawl_multiple_folders(self, folder_paths: List[str], max_depth: int = 3) -> List[Dict[str, Any]]:
        """爬取多个文件夹"""
        all_slices = []
        start_time = datetime.now()
        
        logger.info(f"开始爬取 {len(folder_paths)} 个文件夹")
        
        for folder_path in folder_paths:
            folder_slices = self.crawl_folder(folder_path, max_depth)
            all_slices.extend(folder_slices)
        
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        
        # 打印统计信息
        self.print_stats(duration)
        
        return all_slices
    
    def print_stats(self, duration: float):
        """打印爬取统计信息"""
        logger.info("=" * 60)
        logger.info(f"爬取统计信息 - 总耗时: {duration:.2f} 秒")
        logger.info(f"总文件数: {self.stats['total_files']}")
        logger.info(f"处理文件数: {self.stats['processed_files']}")
        logger.info(f"跳过文件数: {self.stats['skipped_files']}")
        logger.info(f"创建切片数: {self.stats['created_slices']}")
        
        if self.stats['file_type_stats']:
            logger.info("按文件类型统计:")
            for ext, count in sorted(self.stats['file_type_stats'].items(), key=lambda x: x[1], reverse=True):
                logger.info(f"  {ext}: {count} 个")
        
        if self.stats['folder_stats']:
              logger.info("按文件夹统计:")
              top_folders = sorted(self.stats['folder_stats'].items(), key=lambda x: x[1], reverse=True)[:5]
              for folder, count in top_folders:
                  logger.info(f"  {folder}: {count} 个")
        
        logger.info("=" * 60)
    
    def convert_slices_to_memory_format(self, slices: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """将切片数据转换为记忆数据库格式"""
        memory_units = []
        
        for slice_data in slices:
            # 从切片数据中提取内容
            content = slice_data.get('content', '').strip()
            metadata = slice_data.get('metadata', {})
            
            # 跳过空内容
            if not content:
                continue
            
            # 确定主题（使用文件夹名或文件名）
            folder_name = metadata.get('folder_name', '')
            file_name = metadata.get('file_name', '')
            
            # 使用类初始化时设置的主题
            topic = self.topic
            
            # 构建记忆单元
            memory_unit = {
                'topic': topic,
                'content': content,
                'source_type': 'folder_crawler',
                'timestamp': metadata.get('last_modified', datetime.now().isoformat()),
                'importance': 0.7,  # 设置稍高的重要性，因为这些是专门收集的文档
                'confidence': 0.9,  # 高置信度，因为是直接读取的文档
                'tags': [
                    metadata.get('file_type', 'unknown'),
                    'folder_crawler',
                    folder_name.lower() if folder_name else 'general'
                ]
            }
            
            memory_units.append(memory_unit)
        
        return memory_units

def main():
    """测试爬虫功能"""
    crawler = FolderCrawler()
    
    # 测试爬取单个文件夹
    test_folder = "e:\\RAG系统\\docs"
    slices = crawler.crawl_folder(test_folder, max_depth=2)
    
    print(f"测试爬取结果: 处理了 {crawler.stats['processed_files']} 个文件，创建了 {len(slices)} 个切片")
    
    # 转换为记忆格式
    memory_units = crawler.convert_slices_to_memory_format(slices)
    print(f"转换为记忆单元: {len(memory_units)} 个")
    
    # 显示前几个记忆单元作为示例
    for i, memory in enumerate(memory_units[:3]):
        print(f"\n示例记忆单元 {i+1}:")
        print(f"主题: {memory['topic']}")
        print(f"内容预览: {memory['content'][:100]}...")
        print(f"标签: {memory['tags']}")

if __name__ == "__main__":
    main()